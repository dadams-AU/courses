#!/usr/bin/env python3
import os
import glob
from bs4 import BeautifulSoup
from docx import Document

def extract_notes_from_file(file_path):
    """
    Extracts slide content and speaker notes from a reveal.js index.html file.
    Returns a tuple with the presentation title and a list of slide texts.
    """
    print(f"Extracting notes from: {file_path}")
    
    with open(file_path, 'r', encoding='utf-8') as f:
        html_content = f.read()

    soup = BeautifulSoup(html_content, 'html.parser')

    # Extract presentation title or fallback to directory name
    title_tag = soup.find('title')
    title = title_tag.get_text(strip=True) if title_tag else os.path.basename(os.path.dirname(file_path))
    
    print(f"Extracted Title: {title}")

    notes = []
    slides_container = soup.find('div', class_='slides')
    slide_sections = slides_container.find_all('section', recursive=True) if slides_container else soup.find_all('section')

    for index, slide in enumerate(slide_sections, start=1):
        print(f"Processing Slide {index}")
        
        aside_notes = slide.find('aside', class_='notes')
        if aside_notes:
            slide_text = aside_notes.get_text(separator='\n', strip=True)
        else:
            for aside in slide.find_all('aside'):
                aside.decompose()
            slide_text = slide.get_text(separator='\n', strip=True)

        notes.append(f"Slide {index}:\n{slide_text}\n")

    return title, notes

def generate_docx(title, slide_notes, output_docx_path):
    """
    Creates a DOCX file with the extracted presentation notes.
    """
    print(f"Generating DOCX for: {title} -> {output_docx_path}")

    doc = Document()
    doc.add_heading(f"Presentation Title: {title}", level=1)

    for note in slide_notes:
        doc.add_paragraph(note)
        doc.add_paragraph("\n" + "-" * 50 + "\n")  # Separator between slides

    doc.save(output_docx_path)

def main():
    """
    Recursively find all reveal.js index.html files, extract notes, and generate DOCX files.
    """
    print("Scanning for index.html files...")

    index_files = glob.glob('**/index.html', recursive=True)
    if not index_files:
        print("No index.html files found.")
        return

    output_dir = "accessible_docs"
    os.makedirs(output_dir, exist_ok=True)

    for file_path in index_files:
        print(f"\nProcessing: {file_path}")
        try:
            title, slide_notes = extract_notes_from_file(file_path)

            dir_name = os.path.basename(os.path.dirname(file_path))
            docx_filename = f"{dir_name}.docx"
            output_docx_path = os.path.join(output_dir, docx_filename)

            generate_docx(title, slide_notes, output_docx_path)
            print(f"DOCX created: {output_docx_path}\n")

        except Exception as e:
            print(f"Error processing {file_path}: {e}")

if __name__ == '__main__':
    main()

